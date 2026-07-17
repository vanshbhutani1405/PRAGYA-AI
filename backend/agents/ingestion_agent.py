import io
import uuid

import fitz  # PyMuPDF
import openpyxl
from docx import Document as DocxDocument

from services import embedding_service, llm_service, neo4j_service, supabase_service

_TEXT_PDF_MIN_CHARS = 100
_CHUNK_TOKENS = 512
_CHUNK_OVERLAP = 64
_CHARS_PER_TOKEN = 4


def _detect_type(filename: str) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return "pdf"
    if lower.endswith(".xlsx"):
        return "xlsx"
    if lower.endswith(".docx"):
        return "docx"
    if lower.endswith((".png", ".jpg", ".jpeg", ".tiff", ".bmp")):
        return "image"
    return "unknown"


def _parse_pdf(file_bytes: bytes) -> tuple[str, dict | None]:
    """Return (text, entities). Falls back to vision extraction for scanned PDFs."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text = "\n".join(page.get_text() for page in doc)
    if len(text.strip()) >= _TEXT_PDF_MIN_CHARS:
        return text, None
    page = doc.load_page(0)
    pixmap = page.get_pixmap()
    entities = llm_service.vision_extract(pixmap.tobytes("png"))
    return entities["summary"], entities


def _parse_xlsx(file_bytes: bytes) -> str:
    workbook = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    lines = []
    for sheet in workbook.worksheets:
        lines.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                lines.append("\t".join(cells))
    return "\n".join(lines)


def _parse_docx(file_bytes: bytes) -> str:
    document = DocxDocument(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text for c in row.cells if c.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts)


def _chunk_text(text: str) -> list[str]:
    window = _CHUNK_TOKENS * _CHARS_PER_TOKEN
    overlap = _CHUNK_OVERLAP * _CHARS_PER_TOKEN
    step = window - overlap
    chunks = []
    start = 0
    while start < len(text):
        chunk = text[start : start + window]
        if chunk.strip():
            chunks.append(chunk)
        start += step
    return chunks


def _update_graph(entities: dict) -> None:
    for tag in entities.get("equipment_tags", []):
        neo4j_service.upsert_entity_node(tag, "equipment", {})
    for permit in entities.get("permit_ids", []):
        neo4j_service.upsert_entity_node(permit, "permit", {})
    for transition in entities.get("state_transitions", []):
        tag = transition.get("tag")
        if tag:
            neo4j_service.add_state_transition(tag, transition)


def ingest(filename: str, file_bytes: bytes) -> dict:
    """Full ingestion pipeline: parse -> extract -> graph -> chunk -> embed -> index."""
    file_type = _detect_type(filename)

    entities = None
    if file_type == "pdf":
        text, entities = _parse_pdf(file_bytes)
    elif file_type == "xlsx":
        text = _parse_xlsx(file_bytes)
    elif file_type == "docx":
        text = _parse_docx(file_bytes)
    elif file_type == "image":
        entities = llm_service.vision_extract(file_bytes)
        text = entities["summary"]
    else:
        raise ValueError(f"Unsupported file type: {filename}")

    if entities is None:
        entities = llm_service.extract_entities(text)

    _update_graph(entities)

    doc_id = str(uuid.uuid4())
    metadata = {
        "doc_id": doc_id,
        "filename": filename,
        "document_type": entities.get("document_type", ""),
    }

    chunk_texts = _chunk_text(text)
    chunks = [
        {
            "content": chunk,
            "embedding": embedding_service.encode(chunk),
            "chunk_index": index,
        }
        for index, chunk in enumerate(chunk_texts)
    ]
    if chunks:
        supabase_service.index_chunks(chunks, metadata)

    supabase_service.insert_document(
        {
            "id": doc_id,
            "filename": filename,
            "document_type": entities.get("document_type", ""),
            "summary": entities.get("summary", ""),
        }
    )

    return {
        "doc_id": doc_id,
        "filename": filename,
        "chunk_count": len(chunks),
        "entities": entities,
    }
